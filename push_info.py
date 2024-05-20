# -*- coding:utf-8 -*-
import os
import shutil
import docx
import pyperclip
from PIL import ImageGrab
from . import config
from . import common

def data_file_of_today_create():
    if not os.path.exists(config.data_file_path_of_today()):
        shutil.copy(
            config.data_file_path_of_template(),
            config.data_file_path_of_today()
        )

def push_text_of_info(index_of_table=0):
    data_file_of_today_create()
    document=docx.Document(config.data_file_path_of_today())
    table=document.tables[index_of_table]
    cell_col=1
    cell_row=common.count_non_empty_cells_in_column(table,cell_col)
    cell=table.cell(cell_row,cell_col)
    clipboard_text=pyperclip.paste()
    cell.text=clipboard_text
    document.save(config.data_file_path_of_today())
    print(f'{config.data_file_path_of_today()} table[{index_of_table}] index[{cell_row}] 添加文本成功!')

def push_image_of_info(index_of_table=0):
    data_file_of_today_create()
    image_path=0
    image=ImageGrab.grabclipboard()
    if image is None:
        print('剪切板中没有图像数据')
        return False
    else:
        image_path=config.root_dir_path()+'/__clipboard_image__.png'
        image.save(image_path,'PNG')

    document=docx.Document(config.data_file_path_of_today())
    table=document.tables[index_of_table]
    cell_col=1
    cell_row=common.count_non_empty_cells_in_column(table,cell_col)
    if(cell_row==0):
        return False
    cell=table.cell(cell_row-1,cell_col)
    paragraph=cell.add_paragraph()
    paragraph.add_run().add_picture(image_path,width=cell.width)
    document.save(config.data_file_path_of_today())
    if os.path.exists(image_path):
        os.remove(image_path)
    print(f'{config.data_file_path_of_today()} table[{index_of_table}] index[{cell_row-1}] 添加图片成功!')