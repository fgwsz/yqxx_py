# -*- coding:utf-8 -*-
import docx
from yqxx import config
from yqxx import common

def pop_info(index_of_table=0):
    document=docx.Document(config.data_file_path_of_today())
    table=document.tables[index_of_table]
    cell_col=1
    cell_row=common.count_non_empty_cells_in_column(table,cell_col)
    if(cell_row==0):
        exit()
    cell=table.cell(cell_row-1,cell_col)
    # 清空单元格中的所有内容
    for paragraph in cell.paragraphs:
        # 清空段落中的文本
        paragraph.clear()
        # 删除段落中的图片
        for run in paragraph.runs:
            for picture in run.inline_shapes:
                paragraph._p.remove(run._r)
    # 获取单元格中的所有段落
    paragraphs=cell.paragraphs
    # 清空单元格中的所有段落，只保留一个段落
    while len(paragraphs)>1:
        cell._element.remove(paragraphs[-1]._element)
        del paragraphs[-1]
    document.save(config.data_file_path_of_today())
    print(f"{config.data_file_path_of_today()} table[{index_of_table}] index[{cell_row-1}] 清空成功!")
