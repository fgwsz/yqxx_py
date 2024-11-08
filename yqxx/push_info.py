# -*- coding:utf-8 -*-
import os
import docx
import pyperclip
import platform
from yqxx import config
from yqxx import common

def image_grabclipboard_and_save(image_path):
    system_type=platform.system()
    ret=False
    if system_type=='Linux':
        import sys
        from PyQt5.QtWidgets import QApplication
        from PyQt5.QtGui import QClipboard, QImage
        app=QApplication(sys.argv)
        clipboard=app.clipboard()
        # 获取剪贴板中的图像类型
        mime_data=clipboard.mimeData()
        if mime_data.hasImage():
            image=clipboard.image()
            image.save(image_path)
            ret=True
        else:
            print('The clipboard does not contain any images.')
        app.exit()
    else:
        from PIL import ImageGrab
        image=ImageGrab.grabclipboard()
        if image is None:
            print('The clipboard does not contain any images.')
        else:
            image.save(image_path,'PNG')
            ret=True
    return ret

yqxx_white_list=[
    '平原','德城','禹城','夏津','临邑'
]

def push_text_of_info(index_of_table=0):
    common.data_file_of_today_create()
    document=docx.Document(config.data_file_path_of_today())
    table=document.tables[index_of_table]
    total_rows=len(table.rows) #获取表格的总行数
    cell_col=1
    cell_row=common.count_non_empty_cells_in_column(table,cell_col)
    if cell_row==total_rows:
        table.add_row()
        index_cell=table.cell(cell_row,0)
        index_cell.text=f'{cell_row+1}.'
    cell=table.cell(cell_row,cell_col)
    clipboard_text=pyperclip.paste()
    #search white list
    input_flag=True
    for location_item in yqxx_white_list:
        if location_item in clipboard_text:
            print(f'[FAIL][PUSH_TEXT]: white list \'{location_item}\' info!')
            return False

    cell.text=clipboard_text
    document.save(config.data_file_path_of_today())
    print(f'[PASS][PUSH_TEXT]: {config.data_file_path_of_today()} table[{index_of_table}] index[{cell_row}]')
    return True

def push_image_of_info(index_of_table=0):
    common.data_file_of_today_create()
    image_path=config.root_dir_path()+'/__clipboard_image__.png'
    if not image_grabclipboard_and_save(image_path):
        print(f'[FAIL][PUSH_IMAGE]: The clipboard does not contain any images.')
        return False

    document=docx.Document(config.data_file_path_of_today())
    table=document.tables[index_of_table]
    cell_col=1
    cell_row=common.count_non_empty_cells_in_column(table,cell_col)
    if(cell_row==0):
        print(f'[FAIL][PUSH_IMAGE]: no info item text data.')
        return False
    cell=table.cell(cell_row-1,cell_col)
    paragraph=cell.add_paragraph()
    paragraph.add_run().add_picture(image_path,width=cell.width)
    document.save(config.data_file_path_of_today())
    if os.path.exists(image_path):
        os.remove(image_path)
    print(f'[PASS][PUSH_IMAGE]: {config.data_file_path_of_today()} table[{index_of_table}] index[{cell_row-1}]')
    return True
